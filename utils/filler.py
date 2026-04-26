"""
utils/filler.py — Programmatic docx field filler.

Takes a template .docx path, a field schema, and a JSON value mapping,
then writes each value into the correct cell / paragraph.

The LLM is NEVER called here — this is pure python-docx manipulation.
"""
from __future__ import annotations

import copy
import os
import shutil
import tempfile
import time
from typing import Any, Dict, List, Optional


def fill_template(
    template_path: str,
    schema: Dict,
    values: Dict[str, str],
    output_dir: str,
) -> Dict[str, Any]:
    """
    Copy the template .docx and fill it with values from the JSON mapping.

    Args:
        template_path: Path to the original (unmodified) template .docx
        schema:        Template schema dict with "fields" list
        values:        JSON mapping  { "field_0": "...", "field_1": "..." }
        output_dir:    Directory to write the filled file into

    Returns:
        { "status": "ok", "path": ..., "filename": ..., "type": "docx" }
        or
        { "status": "error", "message": ... }
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return {"status": "error", "message": "python-docx not installed. Run: pip install python-docx"}

    if not os.path.isfile(template_path):
        return {"status": "error", "message": f"Template file not found: {template_path}"}

    # ── Copy the template so the original is never modified ──────────────────
    ts = int(time.time())
    base = os.path.splitext(os.path.basename(template_path))[0]
    out_filename = f"filled_{base}_{ts}.docx"
    out_path = os.path.join(output_dir, out_filename)
    shutil.copy2(template_path, out_path)

    try:
        doc = Document(out_path)
        fields = schema.get("fields", [])

        for field in fields:
            fid = field["id"]
            value = values.get(fid, "")

            # "USE_FULL_CONTENT" sentinel — the caller will have resolved this
            # before calling fill_template, but guard here too
            if value == "USE_FULL_CONTENT":
                continue

            loc = field.get("location", {})
            loc_type = loc.get("type", "")

            if loc_type == "table_cell":
                _fill_table_cell(doc, loc, value, field)

            elif loc_type == "paragraph_range":
                _fill_paragraph_range(doc, loc, value, field)

        doc.save(out_path)
        return {
            "status": "ok",
            "path": out_path,
            "filename": out_filename,
            "type": "docx",
        }
    except Exception as e:
        # Clean up failed output
        if os.path.exists(out_path):
            os.remove(out_path)
        return {"status": "error", "message": str(e)}


# ── Table cell filler ─────────────────────────────────────────────────────────

def _fill_table_cell(doc, loc: Dict, value: str, field: Dict):
    """Write value into the specified table cell, preserving font where possible."""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        table_idx = loc["table_index"]
        row_idx   = loc["row_index"]
        cell_idx  = loc["cell_index"]

        if table_idx >= len(doc.tables):
            return
        table = doc.tables[table_idx]

        if row_idx >= len(table.rows):
            return
        row = table.rows[row_idx]

        if cell_idx >= len(row.cells):
            return
        cell = row.cells[cell_idx]

        # Try to copy font/size from neighbour label cell for consistency
        ref_run = _get_neighbour_run(table, row_idx, cell_idx)

        # Clear existing paragraphs in cell
        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ""

        # Write value into first paragraph (preserve paragraph formatting)
        if cell.paragraphs:
            para = cell.paragraphs[0]
            if para.runs:
                run = para.runs[0]
            else:
                run = para.add_run()
            _copy_run_format(run, ref_run)
            run.text = str(value).strip()
        else:
            para = cell.add_paragraph()
            run = para.add_run(str(value).strip())
            _copy_run_format(run, ref_run)

    except Exception:
        # Best-effort: if styling fails, write plain text
        try:
            cell.text = str(value).strip()
        except Exception:
            pass


def _get_neighbour_run(table, row_idx: int, cell_idx: int):
    """Return the first run from an adjacent label cell, for font reference."""
    row = table.rows[row_idx]
    for idx in [cell_idx - 1, cell_idx + 1]:
        if 0 <= idx < len(row.cells):
            neighbour = row.cells[idx]
            for para in neighbour.paragraphs:
                if para.runs:
                    return para.runs[0]
    return None


def _copy_run_format(target_run, ref_run):
    """Copy font properties from ref_run to target_run if ref_run is available."""
    if ref_run is None:
        return
    try:
        from docx.shared import Pt
        if ref_run.font.name:
            target_run.font.name = ref_run.font.name
        if ref_run.font.size:
            target_run.font.size = ref_run.font.size
        if ref_run.font.bold is not None:
            target_run.font.bold = False  # labels are bold; values usually not
        if ref_run.font.color and ref_run.font.color.rgb:
            target_run.font.color.rgb = ref_run.font.color.rgb
    except Exception:
        pass


# ── Paragraph range filler ────────────────────────────────────────────────────

def _fill_paragraph_range(doc, loc: Dict, value: str, field: Dict):
    """
    Replace one or more paragraphs with the provided value text.
    Long-form content is split on double newlines and each chunk becomes a paragraph.
    """
    try:
        start = loc["start_paragraph_index"]
        end   = loc.get("end_paragraph_index", start)

        if start >= len(doc.paragraphs):
            return

        # Get the reference paragraph for style
        ref_para = doc.paragraphs[start]
        ref_style = ref_para.style

        # Split the value into paragraphs
        chunks = [c.strip() for c in str(value).split("\n\n") if c.strip()]
        if not chunks:
            chunks = [str(value).strip()]

        # Replace the first target paragraph with first chunk
        _set_paragraph_text(ref_para, chunks[0])

        # Insert additional paragraphs after the first one if needed
        if len(chunks) > 1:
            insert_after = ref_para
            for chunk in chunks[1:]:
                new_para = _insert_paragraph_after(insert_after, chunk, ref_style)
                insert_after = new_para

    except Exception:
        pass


def _set_paragraph_text(para, text: str):
    """Clear a paragraph and set new text, preserving paragraph style."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _insert_paragraph_after(ref_para, text: str, style=None):
    """Insert a new paragraph immediately after ref_para."""
    from docx.oxml import OxmlElement
    new_para_elem = copy.deepcopy(ref_para._element)
    ref_para._element.addnext(new_para_elem)

    # Get the new paragraph object
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_para_elem, ref_para._parent)

    # Clear runs and set text
    for run in new_para.runs:
        run.text = ""
    if new_para.runs:
        new_para.runs[0].text = text
    else:
        new_para.add_run(text)

    if style:
        try:
            new_para.style = style
        except Exception:
            pass

    return new_para
