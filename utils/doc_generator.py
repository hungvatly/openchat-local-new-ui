"""
OpenChat Local — Document Generator
Creates Word (.docx), PDF (.pdf), and Excel (.xlsx) files from AI-generated content.
Supports both auto-detection from chat and explicit API-driven generation.
"""
import os
import re
import json
import uuid
from typing import Dict, Optional

OUTPUT_DIR = os.path.join("data", "generated")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def generate_docx(title: str, content: str, filename: str = None) -> Dict:
    """Generate a Word document from markdown/text content."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ── Set default style ──
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Calibri"

        # ── Title ──
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph("")  # spacer

        # ── Parse content blocks ──
        # Split on double newlines, but keep code blocks intact
        blocks = _split_blocks(content)

        in_table = False
        table_rows = []

        for block in blocks:
            block = block.strip()
            if not block:
                continue

            # ── Code blocks ──
            code_match = re.match(r'^```(\w*)\n([\s\S]*?)```$', block)
            if code_match:
                lang = code_match.group(1) or "code"
                code = code_match.group(2).strip()
                p = doc.add_paragraph()
                p.style = doc.styles["Normal"]
                run = p.add_run(code)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                # Light gray background via paragraph shading
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F2F2F2')
                shading.set(qn('w:val'), 'clear')
                p.paragraph_format.element.get_or_add_pPr().append(shading)
                continue

            # ── Tables (markdown pipe format) ──
            if "|" in block and block.count("|") >= 2:
                lines = [l.strip() for l in block.split("\n") if l.strip()]
                # Filter separator rows
                data_lines = [l for l in lines if not re.match(r'^[\|\s\-:]+$', l)]
                if len(data_lines) >= 2:
                    rows = []
                    for line in data_lines:
                        cells = [c.strip() for c in line.strip("|").split("|")]
                        rows.append(cells)
                    if rows:
                        cols = max(len(r) for r in rows)
                        table = doc.add_table(rows=len(rows), cols=cols)
                        table.style = "Light Grid Accent 1"
                        for i, row_data in enumerate(rows):
                            for j, cell_val in enumerate(row_data):
                                if j < cols:
                                    table.rows[i].cells[j].text = cell_val
                        doc.add_paragraph("")  # spacer
                        continue

            # ── Headings ──
            if block.startswith("# "):
                doc.add_heading(block[2:], level=1)
            elif block.startswith("## "):
                doc.add_heading(block[3:], level=2)
            elif block.startswith("### "):
                doc.add_heading(block[4:], level=3)
            elif block.startswith("#### "):
                doc.add_heading(block[5:], level=4)
            # ── Bullet lists ──
            elif block.startswith("- ") or block.startswith("* "):
                for line in block.split("\n"):
                    line = line.strip()
                    if line.startswith(("- ", "* ")):
                        p = doc.add_paragraph(style="List Bullet")
                        _add_formatted_text(p, line[2:])
                    elif line:
                        doc.add_paragraph(line)
            # ── Numbered lists ──
            elif re.match(r'^\d+\.\s', block):
                for line in block.split("\n"):
                    line = line.strip()
                    m = re.match(r'^\d+\.\s(.+)', line)
                    if m:
                        p = doc.add_paragraph(style="List Number")
                        _add_formatted_text(p, m.group(1))
                    elif line:
                        doc.add_paragraph(line)
            # ── Blockquotes ──
            elif block.startswith("> "):
                quote_text = "\n".join(l.lstrip("> ") for l in block.split("\n"))
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(quote_text)
                run.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            # ── Horizontal rule ──
            elif block.strip() in ("---", "***", "___"):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("─" * 60)
                run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                run.font.size = Pt(8)
            # ── Normal paragraph ──
            else:
                p = doc.add_paragraph()
                _add_formatted_text(p, block)

        fname = filename or f"{uuid.uuid4().hex[:8]}_{_slug(title)}.docx"
        fpath = os.path.join(OUTPUT_DIR, fname)
        doc.save(fpath)
        return {"status": "ok", "path": fpath, "filename": fname, "type": "docx"}
    except ImportError:
        return {"status": "error", "message": "python-docx not installed"}
    except Exception as e:
        return {"status": "error", "message": str(e)}


def _add_formatted_text(paragraph, text: str):
    """Add text with bold/italic/code inline formatting to a paragraph."""
    from docx.shared import Pt, RGBColor

    # Pattern: **bold**, *italic*, `code`, ***bold italic***
    parts = re.split(r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|`.+?`)', text)
    for part in parts:
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            paragraph.add_run(part)


def _split_blocks(content: str) -> list:
    """Split content into blocks, keeping code fences intact."""
    blocks = []
    current = []
    in_code = False

    for line in content.split("\n"):
        if line.strip().startswith("```") and not in_code:
            # Start of code block — flush current
            if current:
                blocks.extend("\n".join(current).split("\n\n"))
                current = []
            in_code = True
            current.append(line)
        elif line.strip().startswith("```") and in_code:
            # End of code block
            current.append(line)
            blocks.append("\n".join(current))
            current = []
            in_code = False
        elif in_code:
            current.append(line)
        else:
            current.append(line)

    if current:
        blocks.extend("\n".join(current).split("\n\n"))

    return blocks


def generate_pdf(title: str, content: str, filename: str = None) -> Dict:
    """Generate a PDF from markdown/text content with proper formatting."""
    try:
        from fpdf import FPDF
        import sys

        font_family = "Helvetica"
        if sys.platform == "darwin":
            font_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
            font_b = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
            font_i = "/System/Library/Fonts/Supplemental/Arial Italic.ttf"
        elif sys.platform == "win32":
            base = os.environ.get("WINDIR", "C:\\Windows")
            font_path = os.path.join(base, "Fonts", "arial.ttf")
            font_b = os.path.join(base, "Fonts", "arialbd.ttf")
            font_i = os.path.join(base, "Fonts", "ariali.ttf")
        else:
            font_path = "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"
            font_b = "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"
            font_i = "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf"

        class StyledPDF(FPDF):
            def header(self):
                self.set_font(font_family, "I", 8)
                self.set_text_color(150, 150, 150)
                self.cell(0, 8, "OpenChat Local", align="R")
                self.ln(4)

            def footer(self):
                self.set_y(-15)
                self.set_font(font_family, "I", 8)
                self.set_text_color(150, 150, 150)
                self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

        pdf = StyledPDF()
        if os.path.exists(font_path):
            try:
                pdf.add_font("CustomArial", "", font_path)
                if os.path.exists(font_b):
                    pdf.add_font("CustomArial", "B", font_b)
                if os.path.exists(font_i):
                    pdf.add_font("CustomArial", "I", font_i)
                font_family = "CustomArial"
            except Exception:
                pass

        pdf.alias_nb_pages()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()

        # ── Title ──
        pdf.set_font(font_family, "B", 20)
        pdf.set_text_color(30, 30, 30)
        pdf.cell(0, 14, title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(100, 100, 240)
        pdf.set_line_width(0.8)
        pdf.line(10, pdf.get_y(), 80, pdf.get_y())
        pdf.ln(10)

        blocks = _split_blocks(content)

        for block in blocks:
            block = block.strip()
            if not block:
                continue

            # Strip markdown formatting for clean text
            clean = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', block)
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            clean = re.sub(r'`(.+?)`', r'\1', clean)

            # ── Code blocks ──
            code_match = re.match(r'^```(\w*)\n([\s\S]*?)```$', block)
            if code_match:
                code = code_match.group(2).strip()
                pdf.set_fill_color(242, 242, 242)
                pdf.set_font("Courier", "", 9)
                pdf.set_text_color(50, 50, 50)
                # Draw code block with background
                x = pdf.get_x()
                y = pdf.get_y()
                for code_line in code.split("\n"):
                    pdf.cell(0, 5, "  " + code_line, new_x="LMARGIN", new_y="NEXT", fill=True)
                pdf.ln(4)
                pdf.set_text_color(30, 30, 30)
                continue

            # ── Headings ──
            if clean.startswith("# "):
                pdf.set_font(font_family, "B", 18)
                pdf.set_text_color(30, 30, 30)
                pdf.cell(0, 12, clean[2:], new_x="LMARGIN", new_y="NEXT")
                pdf.ln(4)
            elif clean.startswith("## "):
                pdf.set_font(font_family, "B", 15)
                pdf.set_text_color(50, 50, 50)
                pdf.cell(0, 10, clean[3:], new_x="LMARGIN", new_y="NEXT")
                pdf.ln(3)
            elif clean.startswith("### "):
                pdf.set_font(font_family, "B", 13)
                pdf.set_text_color(60, 60, 60)
                pdf.cell(0, 9, clean[4:], new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
            elif clean.startswith("#### "):
                pdf.set_font(font_family, "B", 11)
                pdf.set_text_color(70, 70, 70)
                pdf.cell(0, 8, clean[5:], new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
            # ── Bullet lists ──
            elif clean.startswith("- ") or clean.startswith("* "):
                pdf.set_font(font_family, "", 11)
                pdf.set_text_color(30, 30, 30)
                for line in clean.split("\n"):
                    line = line.strip()
                    if line.startswith(("- ", "* ")):
                        pdf.cell(8)  # indent
                        pdf.cell(5, 6, "\u2022")
                        pdf.multi_cell(0, 6, line[2:])
                    elif line:
                        pdf.multi_cell(0, 6, line)
                pdf.ln(2)
            # ── Numbered lists ──
            elif re.match(r'^\d+\.\s', clean):
                pdf.set_font(font_family, "", 11)
                pdf.set_text_color(30, 30, 30)
                for line in clean.split("\n"):
                    line = line.strip()
                    m = re.match(r'^(\d+)\.\s(.+)', line)
                    if m:
                        pdf.cell(8)
                        pdf.cell(8, 6, f"{m.group(1)}.")
                        pdf.multi_cell(0, 6, m.group(2))
                    elif line:
                        pdf.multi_cell(0, 6, line)
                pdf.ln(2)
            # ── Blockquotes ──
            elif clean.startswith("> "):
                quote_text = "\n".join(l.lstrip("> ") for l in clean.split("\n"))
                pdf.set_font(font_family, "I", 11)
                pdf.set_text_color(100, 100, 100)
                pdf.set_draw_color(180, 180, 180)
                x = pdf.get_x()
                pdf.line(x + 4, pdf.get_y(), x + 4, pdf.get_y() + 8)
                pdf.cell(10)
                pdf.multi_cell(0, 6, quote_text)
                pdf.ln(3)
                pdf.set_text_color(30, 30, 30)
            # ── Horizontal rule ──
            elif clean.strip() in ("---", "***", "___"):
                pdf.set_draw_color(200, 200, 200)
                y = pdf.get_y()
                pdf.line(10, y, 200, y)
                pdf.ln(6)
            # ── Tables ──
            elif "|" in clean and clean.count("|") >= 2:
                lines = [l.strip() for l in clean.split("\n") if l.strip()]
                data_lines = [l for l in lines if not re.match(r'^[\|\s\-:]+$', l)]
                if data_lines:
                    rows = [[c.strip() for c in l.strip("|").split("|")] for l in data_lines]
                    cols = max(len(r) for r in rows)
                    col_w = (pdf.w - 20) / cols
                    for i, row in enumerate(rows):
                        if i == 0:
                            pdf.set_font(font_family, "B", 10)
                            pdf.set_fill_color(68, 114, 196)
                            pdf.set_text_color(255, 255, 255)
                        else:
                            pdf.set_font(font_family, "", 10)
                            pdf.set_fill_color(245, 245, 245) if i % 2 == 0 else pdf.set_fill_color(255, 255, 255)
                            pdf.set_text_color(30, 30, 30)
                        for j, cell in enumerate(row):
                            if j < cols:
                                pdf.cell(col_w, 7, cell[:30], border=1, fill=True)
                        pdf.ln()
                    pdf.set_text_color(30, 30, 30)
                    pdf.ln(4)
            # ── Normal paragraph ──
            else:
                pdf.set_font(font_family, "", 11)
                pdf.set_text_color(30, 30, 30)
                pdf.multi_cell(0, 6, clean)
                pdf.ln(3)

        fname = filename or f"{uuid.uuid4().hex[:8]}_{_slug(title)}.pdf"
        fpath = os.path.join(OUTPUT_DIR, fname)
        pdf.output(fpath)
        return {"status": "ok", "path": fpath, "filename": fname, "type": "pdf"}
    except ImportError:
        return {"status": "error", "message": "fpdf2 not installed. Run: pip install fpdf2"}
    except Exception as e:
        return {"status": "error", "message": str(e)}


def generate_xlsx(title: str, content: str, filename: str = None) -> Dict:
    """Generate an Excel file from tabular content."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        wb = Workbook()
        ws = wb.active
        ws.title = title[:31]

        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, size=11, color="FFFFFF")
        alt_fill = PatternFill(start_color="F2F6FC", end_color="F2F6FC", fill_type="solid")
        thin_border = Border(
            left=Side(style="thin", color="DDDDDD"),
            right=Side(style="thin", color="DDDDDD"),
            top=Side(style="thin", color="DDDDDD"),
            bottom=Side(style="thin", color="DDDDDD"),
        )

        lines = [l.strip() for l in content.strip().split("\n") if l.strip()]

        # Try to detect table format (pipe-separated or CSV-like)
        if any("|" in l for l in lines):
            # Markdown table
            row_num = 1
            for line in lines:
                if set(line.replace("|", "").strip()) <= set("- :"):
                    continue  # skip separator row
                cells = [c.strip() for c in line.strip("|").split("|")]
                for col, val in enumerate(cells, 1):
                    cell = ws.cell(row=row_num, column=col, value=val)
                    cell.border = thin_border
                    if row_num == 1:
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.alignment = Alignment(horizontal="center")
                    elif row_num % 2 == 0:
                        cell.fill = alt_fill
                row_num += 1
        elif any("," in l for l in lines):
            # CSV-like
            import csv
            import io
            reader = csv.reader(io.StringIO(content))
            for row_num, row in enumerate(reader, 1):
                for col, val in enumerate(row, 1):
                    cell = ws.cell(row=row_num, column=col, value=val.strip())
                    cell.border = thin_border
                    if row_num == 1:
                        cell.font = header_font
                        cell.fill = header_fill
                    elif row_num % 2 == 0:
                        cell.fill = alt_fill
        else:
            # Just dump lines into column A
            ws.cell(row=1, column=1, value=title).font = Font(bold=True, size=12)
            for i, line in enumerate(lines, 2):
                ws.cell(row=i, column=1, value=line)

        # Auto-width columns
        for col in ws.columns:
            max_len = 0
            col_letter = col[0].column_letter
            for cell in col:
                if cell.value:
                    max_len = max(max_len, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = min(max_len + 4, 50)

        fname = filename or f"{uuid.uuid4().hex[:8]}_{_slug(title)}.xlsx"
        fpath = os.path.join(OUTPUT_DIR, fname)
        wb.save(fpath)
        return {"status": "ok", "path": fpath, "filename": fname, "type": "xlsx"}
    except ImportError:
        return {"status": "error", "message": "openpyxl not installed. Run: pip install openpyxl"}
    except Exception as e:
        return {"status": "error", "message": str(e)}


# ── Detection keywords (broader set) ──────────────────────────────────────────

_DOCX_KEYWORDS = [
    "create a word", "write a docx", "make a word doc", "generate a document",
    "create a report", "write a report", "save as docx", "export as word",
    "make a document", "create document", "generate report", "write document",
    "draft a report", "draft a document", "word file", "docx file",
    "create a letter", "write a letter", "draft a letter",
]

_PDF_KEYWORDS = [
    "create a pdf", "make a pdf", "generate a pdf", "export as pdf",
    "save as pdf", "pdf file", "write a pdf", "create pdf",
    "export to pdf", "convert to pdf", "make pdf",
]

_XLSX_KEYWORDS = [
    "create a spreadsheet", "make an excel", "create an xlsx", "generate a table",
    "create a csv", "make a spreadsheet", "excel file", "xlsx file",
    "create spreadsheet", "generate spreadsheet", "export as excel",
    "save as excel", "export as xlsx", "make excel",
]


def detect_and_generate(ai_response: str, user_message: str) -> Optional[Dict]:
    """Check if AI response contains content meant for a file and generate it."""
    msg_lower = user_message.lower()

    # Detect document creation intent from the user's message
    if any(kw in msg_lower for kw in _DOCX_KEYWORDS):
        title = _extract_title(user_message, ai_response)
        return generate_docx(title, ai_response)

    if any(kw in msg_lower for kw in _PDF_KEYWORDS):
        title = _extract_title(user_message, ai_response)
        return generate_pdf(title, ai_response)

    if any(kw in msg_lower for kw in _XLSX_KEYWORDS):
        title = _extract_title(user_message, ai_response)
        return generate_xlsx(title, ai_response)

    return None


def _slug(text: str) -> str:
    slug = re.sub(r"[^\w\s-]", "", text.lower())
    return re.sub(r"[\s-]+", "_", slug).strip("_")[:40]


def _extract_title(user_msg: str, ai_response: str) -> str:
    """Try to extract a sensible title."""
    all_keywords = _DOCX_KEYWORDS + _PDF_KEYWORDS + _XLSX_KEYWORDS
    for prefix in ["create a ", "make a ", "generate a ", "write a ", "draft a "]:
        if prefix in user_msg.lower():
            after = user_msg.lower().split(prefix, 1)[1]
            # remove format words
            for fmt in ["word doc", "docx", "pdf", "spreadsheet", "excel", "xlsx",
                        "document", "report", "table", "letter", "file", "about", "on", "for"]:
                after = after.replace(fmt, "").strip()
            if after and len(after) > 3:
                return after.strip(" .,!?")[:60].title()
    # Fallback: first heading in AI response, or first line
    for line in ai_response.strip().split("\n"):
        line = line.strip()
        if line.startswith("# "):
            return line[2:][:60]
    first_line = ai_response.strip().split("\n")[0]
    clean = re.sub(r"^#+\s*", "", first_line)
    return clean[:60] if clean else "Document"
