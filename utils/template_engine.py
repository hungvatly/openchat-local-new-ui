"""
utils/template_engine.py — Template store + fill pipeline.

Two-phase architecture:
  Phase 1 (upload-time): scan template structure → save schema JSON + original .docx
  Phase 2 (fill-time):   extract content text → LLM returns JSON map → filler writes docx

The LLM NEVER generates code. Its only job is returning a JSON key→value mapping.
All document construction is done by utils/filler.py using python-docx.
"""
from __future__ import annotations

import json
import os
import re
import shutil
import uuid
from typing import Any, Dict, List, Optional

# ── Directory layout ───────────────────────────────────────────────────────────

_BASE    = os.path.join(os.path.dirname(__file__), "..", "data")
TMPL_DIR = os.path.join(_BASE, "templates")
OUT_DIR  = os.path.join(_BASE, "outputs")

os.makedirs(TMPL_DIR, exist_ok=True)
os.makedirs(OUT_DIR,  exist_ok=True)

# Files path is also served at /files/... by main.py
OUTPUT_DIR = OUT_DIR


# ── Phase 1: Template upload & scan ───────────────────────────────────────────

def save_template(src_path: str, original_filename: str) -> Dict:
    """
    Store the template file and its field schema.

    Returns the full template record dict.
    """
    from utils.scanner import scan_template

    template_id = uuid.uuid4().hex[:12]
    tmpl_dir = os.path.join(TMPL_DIR, template_id)
    os.makedirs(tmpl_dir, exist_ok=True)

    ext = os.path.splitext(original_filename)[1].lower()
    stored_filename = f"original{ext}"
    stored_path = os.path.join(tmpl_dir, stored_filename)
    shutil.copy2(src_path, stored_path)

    # Scan structure (programmatic, not LLM)
    structure = scan_template(stored_path)

    record = {
        "id":         template_id,
        "name":       original_filename,
        "filename":   stored_filename,
        "extension":  ext,
        "path":       stored_path,
        "fields":     structure.get("fields", []),
        "structure":  structure,
    }

    schema_path = os.path.join(tmpl_dir, "schema.json")
    with open(schema_path, "w", encoding="utf-8") as f:
        json.dump(record, f, ensure_ascii=False, indent=2)

    return record


def get_template(template_id: str) -> Optional[Dict]:
    """Load a stored template record by ID."""
    schema_path = os.path.join(TMPL_DIR, template_id, "schema.json")
    if not os.path.isfile(schema_path):
        return None
    with open(schema_path, "r", encoding="utf-8") as f:
        return json.load(f)


def list_templates() -> List[Dict]:
    """Return summary list of all saved templates."""
    results = []
    if not os.path.isdir(TMPL_DIR):
        return results
    for tid in os.listdir(TMPL_DIR):
        schema_path = os.path.join(TMPL_DIR, tid, "schema.json")
        if os.path.isfile(schema_path):
            try:
                with open(schema_path, "r", encoding="utf-8") as f:
                    record = json.load(f)
                results.append({
                    "id":     record["id"],
                    "name":   record["name"],
                    "fields": len(record.get("fields", [])),
                    "extension": record.get("extension", ""),
                })
            except Exception:
                pass
    return sorted(results, key=lambda r: r["name"])


def delete_template(template_id: str) -> bool:
    """Delete a template directory and all its files."""
    tmpl_dir = os.path.join(TMPL_DIR, template_id)
    if os.path.isdir(tmpl_dir):
        shutil.rmtree(tmpl_dir)
        return True
    return False


# ── Phase 2: Fill pipeline ────────────────────────────────────────────────────

_SYSTEM_PROMPT = (
    "You are a form-filling assistant. "
    "You receive a list of template fields and a content document. "
    "Your ONLY job is to extract the correct value for each field from the content. "
    "Respond with ONLY a valid JSON object — no markdown fences, no explanation, "
    "no extra text before or after the JSON."
)

_SHORT_FIELD_NOTE = "(short value: name, date, number, or phrase)"
_LONG_FIELD_NOTE  = (
    "(long text: return the COMPLETE text — do NOT summarise or truncate. "
    "If the content is very long, return the signal USE_FULL_CONTENT instead.)"
)


async def run_fill_pipeline(
    template_id: str,
    content_text: str,
    model: Optional[str] = None,
) -> Dict:
    """
    Full fill pipeline:
      1. Load template schema
      2. Build LLM prompt
      3. Call LLM via engine-agnostic bridge → get JSON map
      4. Resolve USE_FULL_CONTENT sentinels
      5. Call filler.fill_template() to produce the .docx
      6. Return { status, url, filename, type }

    This function is async because the LLM call is async.
    """
    from utils.llm_bridge import extract_json_from_llm
    from utils.filler import fill_template

    template = get_template(template_id)
    if not template:
        return {"status": "error", "error": "Template not found"}

    fields: List[Dict] = template.get("fields", [])
    if not fields:
        return {"status": "error", "error": "Template has no detected fields. Try re-uploading."}

    # ── Step 2: Build prompt ──────────────────────────────────────────────────
    field_lines = []
    for f in fields:
        note = _LONG_FIELD_NOTE if f.get("type") == "long" else _SHORT_FIELD_NOTE
        field_lines.append(f"- {f['id']} ({f['label']}) {note}")

    user_prompt = (
        "Here are the fields to fill:\n\n"
        + "\n".join(field_lines)
        + "\n\nCONTENT DOCUMENT:\n---\n"
        + _truncate_content(content_text)
        + "\n---\n\n"
        "Return a JSON object mapping each field ID to its extracted value.\n"
        "Example format:\n"
        + json.dumps({f["id"]: "value here" for f in fields[:2]}, indent=2)
    )

    # ── Step 3: Call LLM ──────────────────────────────────────────────────────
    values: Dict[str, str] = await extract_json_from_llm(
        system_prompt=_SYSTEM_PROMPT,
        user_prompt=user_prompt,
        model=model,
    )

    if not values:
        return {"status": "error", "error": "LLM did not return a parseable JSON mapping."}

    # ── Step 4: Resolve USE_FULL_CONTENT sentinels ────────────────────────────
    for fid, val in values.items():
        if str(val).strip().upper() == "USE_FULL_CONTENT":
            values[fid] = content_text  # use the full raw content text

    # ── Step 5: Fill the document ─────────────────────────────────────────────
    template_path = template["path"]
    result = fill_template(
        template_path=template_path,
        schema=template,
        values=values,
        output_dir=OUT_DIR,
    )

    if result.get("status") == "ok":
        result["url"] = f"/files/{result['filename']}"

    return result


def _truncate_content(text: str, max_chars: int = 6000) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + f"\n\n[content truncated at {max_chars} chars — {len(text)-max_chars} more chars available]"


# ── Legacy compat shims (kept so existing /api/templates/* routes still compile) ──

def build_fill_prompt(template: Dict, user_instructions: str) -> str:
    """Legacy prompt builder for the instruction-based fill path (AI Fill tab)."""
    fields = template.get("fields", [])
    raw = template.get("structure", {}).get("raw_text", "")
    lines = [
        "You are filling in a document template.",
        "Generate a COMPLETE document following the exact same structure with the fields filled in.",
        "",
        "Fields detected in template:",
    ]
    for f in fields:
        lines.append(f"  - {f['label']} ({f.get('type','short')})")
    lines += ["", "Template text:", raw[:3000], "", "User instructions:", user_instructions, "",
              "Output the filled document as Markdown."]
    return "\n".join(lines)


def generate_from_template(template_id: str, ai_content: str, output_format: str = None) -> Dict:
    """Legacy renderer for the instruction-based fill path (AI Fill tab)."""
    template = get_template(template_id)
    if not template:
        return {"status": "error", "message": "Template not found"}

    ext = output_format or template.get("extension", ".docx")
    title = re.sub(r"\.[^.]+$", "", template["name"])
    import uuid as _uuid
    fname = f"{_uuid.uuid4().hex[:8]}_filled_{re.sub(r'[^a-zA-Z0-9]','_',title)}{ext}"

    if ext == ".docx":
        return _legacy_generate_docx(template, ai_content, fname)
    else:
        from utils.doc_generator import generate_pdf
        return generate_pdf(f"Filled: {title}", ai_content, fname)


def _legacy_generate_docx(template: Dict, ai_content: str, filename: str) -> Dict:
    """Simple Markdown → .docx renderer used by the legacy AI Fill tab."""
    try:
        from docx import Document
        from docx.shared import Pt
        import re as _re

        tmpl_path = template.get("path", "")
        if os.path.exists(tmpl_path) and tmpl_path.endswith(".docx"):
            try:
                doc = Document(tmpl_path)
                for para in doc.paragraphs:
                    para.clear()
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            cell.text = ""
            except Exception:
                doc = Document()
        else:
            doc = Document()

        for block in ai_content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("# "):
                doc.add_heading(block[2:], 1)
            elif block.startswith("## "):
                doc.add_heading(block[3:], 2)
            elif block.startswith("### "):
                doc.add_heading(block[4:], 3)
            elif "|" in block and block.count("|") >= 2:
                lines = [l.strip() for l in block.split("\n") if l.strip()]
                rows = []
                for line in lines:
                    if set(line.replace("|", "").strip()) <= set("- :"):
                        continue
                    rows.append([c.strip() for c in line.strip("|").split("|")])
                if rows:
                    cols = max(len(r) for r in rows)
                    tbl = doc.add_table(rows=len(rows), cols=cols)
                    tbl.style = "Table Grid"
                    for i, r in enumerate(rows):
                        for j, v in enumerate(r):
                            if j < cols:
                                tbl.rows[i].cells[j].text = v
            elif block.startswith(("- ", "* ")):
                for line in block.split("\n"):
                    line = line.strip()
                    if line.startswith(("- ", "* ")):
                        doc.add_paragraph(line[2:], style="List Bullet")
            else:
                clean = _re.sub(r"\*\*(.+?)\*\*", r"\1", block)
                clean = _re.sub(r"\*(.+?)\*", r"\1", clean)
                doc.add_paragraph(clean)

        fpath = os.path.join(OUT_DIR, filename)
        doc.save(fpath)
        return {"status": "ok", "path": fpath, "filename": filename, "type": "docx"}
    except Exception as e:
        return {"status": "error", "message": str(e)}
