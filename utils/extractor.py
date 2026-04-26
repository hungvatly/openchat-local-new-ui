"""
utils/extractor.py — Extract text content from .docx, .pdf, .txt, and .md files.
Used by the template-fill pipeline to read content source files.
"""
from __future__ import annotations
import os
from typing import List


def extract_text(path: str, max_chars: int = 12000) -> str:
    """
    Extract readable text from a file.  Returns plain text with paragraph
    breaks preserved.  Truncates to max_chars to stay within model context.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        text = _from_docx(path)
    elif ext == ".pdf":
        text = _from_pdf(path)
    elif ext in (".txt", ".md", ".csv"):
        text = _from_text(path)
    else:
        # Attempt as plain text
        text = _from_text(path)

    if len(text) > max_chars:
        text = text[:max_chars] + f"\n\n[... content truncated at {max_chars} chars ...]"
    return text


def _from_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_parts = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_parts:
                    parts.append(" | ".join(row_parts))
        return "\n\n".join(parts)
    except ImportError:
        return "(python-docx not installed)"
    except Exception as e:
        return f"(Error reading .docx: {e})"


def _from_pdf(path: str) -> str:
    # Try pdfplumber first (better table support)
    try:
        import pdfplumber
        parts: List[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(text.strip())
        return "\n\n".join(parts)
    except ImportError:
        pass
    # Fallback to pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        return "\n\n".join(
            (p.extract_text() or "").strip()
            for p in reader.pages
            if (p.extract_text() or "").strip()
        )
    except ImportError:
        return "(No PDF library: install pdfplumber)"
    except Exception as e:
        return f"(Error reading PDF: {e})"


def _from_text(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        return f"(Error reading file: {e})"
