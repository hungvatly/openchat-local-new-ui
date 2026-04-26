"""
utils/scanner.py — Programmatic template structure scanner.

Walks a .docx file using python-docx and detects fillable fields:
  - Table cells that sit next to label cells (e.g. "Name:" → adjacent empty cell)
  - Paragraph-level body sections that are empty / placeholder text

Returns a JSON-serialisable schema dict that is stored beside the template file.
"""
from __future__ import annotations

import re
import uuid
from typing import Any, Dict, List, Optional


# ── Placeholder detection ──────────────────────────────────────────────────────

_PLACEHOLDER_RE = re.compile(
    r"^\s*$"                     # completely empty
    r"|^[_\.\u2026\-\s]+$"       # underscores / dots / dashes only
    r"|^\[.*\]$"                  # [bracket placeholder]
    r"|^<.*>$"                    # <angle bracket placeholder>
    r"|\u00a0+"                   # non-breaking spaces
    r"|^\.{3,}$",                 # ellipsis
    re.UNICODE
)

_LABEL_RE = re.compile(r".{2,}")   # any cell with ≥ 2 chars that looks like a label


def _is_placeholder(text: str) -> bool:
    return bool(_PLACEHOLDER_RE.match(text.strip()))


def _is_label(text: str) -> bool:
    t = text.strip()
    return bool(t) and bool(_LABEL_RE.match(t)) and not _is_placeholder(t)


# ── Main scanner ───────────────────────────────────────────────────────────────

def scan_template(path: str) -> Dict[str, Any]:
    """
    Scan a .docx template and return a field schema dict.

    Returns:
        {
            "fields": [ { "id", "label", "type", "location" }, ... ],
            "raw_text": "..."          # full extracted text for context
        }
    """
    if not path.lower().endswith(".docx"):
        return {"fields": [], "raw_text": _scan_non_docx(path)}

    try:
        from docx import Document
    except ImportError:
        return {"fields": [], "raw_text": "(python-docx not installed)"}

    doc = Document(path)
    fields: List[Dict] = []
    raw_lines: List[str] = []
    field_counter = 0

    # ── Scan tables ────────────────────────────────────────────────────────────
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            cells = row.cells
            raw_lines.append(" | ".join(c.text.strip() for c in cells))

            for cell_idx, cell in enumerate(cells):
                cell_text = cell.text.strip()

                if not _is_placeholder(cell_text):
                    continue  # not a fillable cell

                # Look for a label in adjacent cells (left or right)
                label = _find_label_for_cell(cells, cell_idx)
                if label is None:
                    # Look above (previous row, same column)
                    label = _find_label_above(table, row_idx, cell_idx)

                if label:
                    fid = f"field_{field_counter}"
                    field_counter += 1
                    fields.append({
                        "id": fid,
                        "label": label,
                        "type": "short",
                        "location": {
                            "type": "table_cell",
                            "table_index": table_idx,
                            "row_index": row_idx,
                            "cell_index": cell_idx,
                        },
                    })

    # ── Scan paragraphs ────────────────────────────────────────────────────────
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        raw_lines.append(text)

        # Detect long-form body sections: empty paragraphs that follow a heading
        if _is_placeholder(text):
            # Check if the preceding paragraph looks like a heading or section label
            label = _preceding_heading(doc.paragraphs, para_idx)
            if label:
                # Only add if we don't already have a field for this area
                if not any(
                    f["location"]["type"] == "paragraph_range"
                    and f["location"]["start_paragraph_index"] == para_idx
                    for f in fields
                ):
                    fid = f"field_{field_counter}"
                    field_counter += 1
                    fields.append({
                        "id": fid,
                        "label": label,
                        "type": "long",
                        "location": {
                            "type": "paragraph_range",
                            "start_paragraph_index": para_idx,
                            "end_paragraph_index": para_idx,
                        },
                    })

    return {
        "fields": fields,
        "raw_text": "\n".join(raw_lines),
    }


def _find_label_for_cell(cells, cell_idx: int) -> Optional[str]:
    """Return label text from the cell immediately to the left, or right, of cell_idx."""
    # Try left
    if cell_idx > 0:
        left = cells[cell_idx - 1].text.strip()
        if _is_label(left):
            return left.rstrip(":").strip()
    # Try right
    if cell_idx < len(cells) - 1:
        right = cells[cell_idx + 1].text.strip()
        if _is_label(right):
            return right.rstrip(":").strip()
    # Try left-left (two-column skip — label | value | label | value layout)
    if cell_idx > 1:
        ll = cells[cell_idx - 1].text.strip()  # already checked, but try skipping
        if not _is_label(ll) and cell_idx > 1:
            ll2 = cells[cell_idx - 2].text.strip()
            if _is_label(ll2):
                return ll2.rstrip(":").strip()
    return None


def _find_label_above(table, row_idx: int, cell_idx: int) -> Optional[str]:
    """Return label from the same column in the previous row, if it's a label."""
    if row_idx == 0:
        return None
    prev_row = table.rows[row_idx - 1]
    if cell_idx < len(prev_row.cells):
        text = prev_row.cells[cell_idx].text.strip()
        if _is_label(text):
            return text.rstrip(":").strip()
    return None


def _preceding_heading(paragraphs, para_idx: int) -> Optional[str]:
    """Return the text of the nearest preceding heading/label paragraph."""
    for i in range(para_idx - 1, max(para_idx - 6, -1), -1):
        p = paragraphs[i]
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ""
        if "Heading" in style or (len(text) < 80 and not _is_placeholder(text)):
            return text.rstrip(":").strip()
    return None


def _scan_non_docx(path: str) -> str:
    """Fallback for non-.docx files: just extract raw text."""
    try:
        from utils.extractor import extract_text
        return extract_text(path)
    except Exception:
        return ""
